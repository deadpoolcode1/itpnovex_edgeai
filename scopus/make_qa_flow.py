#!/usr/bin/env python3
"""Generate Scopus_QA_Flow.docx — the short, copy-paste QA flow.

This is the *short* document. `Scopus_Tester_Manual.docx` is the long one: it
explains what every step proves, carries the troubleshooting for each, and
covers the parts a tester only reaches when something is wrong. This one is
the flow a QA engineer repeats — two parts, one command per line, nothing to
decide:

    Part 1   the product on the cable, receiving on this PC
    Part 2   the same product over the mobile network, driven from and
             received by the customer's server, with the MQTT command channel

It is generated rather than hand-edited for the same reason as the manual: the
commands stay in one place with the suite, and a change to the procedure is a
reviewable diff instead of a binary blob.

    python3 scopus/make_qa_flow.py

Every command in here was run, in this order, on the bench below, and every
"Expected" block is what came back on 2026-08-17. Where a value changes run to
run — counters, timestamps, the size of a JPEG — the shape is what matters and
the document says so.
"""
import datetime
import pathlib
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

OUT = pathlib.Path(__file__).resolve().parent / "Scopus_QA_Flow.docx"

MONO = "Consolas"
GREY = RGBColor(0x44, 0x44, 0x44)
RED = RGBColor(0xB0, 0x00, 0x00)
GREEN = RGBColor(0x1E, 0x6B, 0x2E)
BLUE = RGBColor(0x1F, 0x3D, 0x7A)

REPO = "~/work/itpnovex/edgeai"
CD = f"cd {REPO} && "

# The machine the two devices are plugged into. Naming it matters: these
# commands only work there, and running them on another PC gives "no such
# device" for reasons no troubleshooting table can explain.
BENCH_HOST = "T7ARYZ0009769Z2"
BENCH_ADDR = "100.115.215.6"
BENCH_SSH_PORT = 4322
BENCH_USER = "user"

# The customer's server. This address is the office public IP, and the
# receiver, the broker and the certificates are already installed on it —
# nothing in this document installs anything.
# ScopusQA #11: the receiver address is site data. It is still printed into
# the generated document — the tester needs it — but it is read from
# scopus/bench.ini (untracked) rather than committed here.
import os as _os, sys as _sys
_sys.path.insert(0, _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "lib"))
from settings import S as _S
SRV = _S.require("server", "host")
SRV_HTTP_PORT = 8991          # ITP's receiver: /upload and /notify
MQTT_PORT = 5912              # mosquitto, TLS + client certificate
CERTS = "/opt/sdvr-server/certs"
UPLOAD_DIR = "/home/user/sdvr-uploads"
APN = "internet"

# The unit's name on the command channel. It is the modem's IMEI, and Step 2
# of Part 2 reads it back off the device rather than trusting this constant —
# a wrong id here would subscribe to a topic nobody publishes on, which looks
# exactly like a unit that is switched off.
IMEI = "359779080290964"

MQTT_AUTH = (f"--cafile {CERTS}/ca.crt \\\n"
             f"  --cert {CERTS}/client.crt --key {CERTS}/client.key")


def code(doc, text, size=9):
    """A command / output block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(size)
    return p


def cmd(doc, text):
    """A line the tester types. Bold so it stands out from expected output."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = BLUE
    return p


def note(doc, text, warn=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RED if warn else GREY
    return p


def expected(doc, text="Expected:"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(2)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    if widths:
        # Three things have to agree or the widths are ignored: autofit off,
        # every cell's own width, and w:tblGrid — which is the one that
        # actually governs under a fixed layout, and which python-docx leaves
        # at equal columns however many cell widths you set.
        t.autofit = False
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
        grid = t._tbl.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            "tblGrid")
        if grid is not None:
            for col, w in zip(grid, widths):
                col.set("{http://schemas.openxmlformats.org/wordprocessingml/"
                        "2006/main}w", str(int(w * 1440)))
    return t


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    doc.add_heading("Scopus — QA Flow", 0)
    p = doc.add_paragraph()
    p.add_run("Two runs of the same product. First on the cable, then over "
              "the mobile network through the customer's server — including "
              "driving the unit from that server.").bold = True
    doc.add_paragraph(
        f"Generated {datetime.date.today().isoformat()} from "
        f"scopus/make_qa_flow.py. Every command below was run in this order "
        f"on {BENCH_HOST}, and every “Expected” block is what came back. "
        f"Counters, timestamps and file sizes change run to run; the shape "
        f"and the marked values do not.")
    note(doc,
         "This is the short document — the flow, and nothing else. "
         "Scopus_Tester_Manual.docx is the long one: what each step proves, "
         "the full troubleshooting, the cellular relay, and the parts you "
         "only reach when something is wrong. Reach for it when a step here "
         "fails and the table at the end does not cover it.")

    # ── Before you start ───────────────────────────────────────────────
    doc.add_heading("1. Before you start", level=1)
    doc.add_paragraph(
        "Everything in Part 1 runs on the PC the two devices are plugged "
        "into. Part 2 splits: the device is configured once from that same "
        "PC, and everything after that is typed on the customer's server.")
    table(doc, ["What", "Where"],
          [["Bench PC (camera + modem plugged in)",
            f"{BENCH_HOST} — ssh -p {BENCH_SSH_PORT} "
            f"{BENCH_USER}@{BENCH_ADDR}"],
           ["Customer's server (receiver + MQTT broker)",
            f"{SRV} — receiver on TCP {SRV_HTTP_PORT}, broker on TCP "
            f"{MQTT_PORT}"],
           ["Repository on the bench", REPO],
           ["Photos and events land in", UPLOAD_DIR]])
    note(doc,
         f"On this bench the two are the same machine — {SRV} is its own "
         f"office public address — so Part 2 is typed in the same terminal. "
         f"At a real site they are two machines, and only the commands that "
         f"start with mosquitto_ move to the server.")

    doc.add_paragraph()
    doc.add_paragraph("You need two terminal windows, and a third if you "
                      "want to watch the camera:")
    table(doc, ["Window", "What it does"],
          [["A", "Left running. It shows what arrives — the receiving server "
                 "in Part 1, the unit's command channel in Part 2."],
           ["B", "Everything you type."],
           ["C (optional)", "Live picture from the camera. Must be opened on "
                            "the bench PC's own screen, not over SSH."]])
    for a in ["Lines in bold blue are what you type. Copy the whole line, "
              "including the part before &&.",
              "Everything else in the boxes is what the computer prints back.",
              "Nothing here changes the product's software.",
              "About 15 minutes for Part 1, 10 for Part 2."]:
        doc.add_paragraph(a, style="List Bullet")

    # ══ PART 1 ═════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("2. Part 1 — Local live test (on the cable)",
                    level=1)
    doc.add_paragraph(
        "The camera detects, the modem sends, and this PC receives — over "
        "the USB cable. Nothing leaves the room. Run this first even when "
        "cellular is what you were asked to test: it uses the same camera, "
        "the same modem and the same internal cable, so a failure here tells "
        "you where the fault is, which the network test cannot.")

    doc.add_heading("Step 1 — Check the bench (Window B)", level=2)
    cmd(doc, f"{CD}python3 scopus/preflight.py")
    expected(doc, "Twelve checks, all PASS, ending in:")
    code(doc, "  [PASS] Camera can reach the modem  —  internal cable carries "
              "commands\n"
              "  [PASS] Test image present  —  .../edgeai/images/3_people.jpg\n"
              "\n"
              "READY — the bench is fit to test on.")
    note(doc, "A failure here is a bench fault, not a product fault. Each one "
              "prints what to do about it — do that first. Chasing a bench "
              "fault as a product fault is where the hours go.", warn=True)

    doc.add_heading("Step 2 — Start the receiving server (Window A)", level=2)
    doc.add_paragraph("This window plays the part of the customer's server. "
                      "Start it and leave it visible.")
    cmd(doc, f"{CD}python3 scopus/test_server.py --http-port 8080 "
             f"--udp-port 9999 \\\n"
             f"  --dir ~/scopus-received --from-modem 192.168.2.2 --fresh")
    expected(doc)
    code(doc, "Scopus test server\n"
              "  receiving into /home/user/scopus-received\n"
              "[09:45:48] listening      UDP  0.0.0.0:9999  (notifications)\n"
              "[09:45:48] listening      HTTP 0.0.0.0:8080 (photo uploads)")

    doc.add_heading("Step 3 — Point the device at this PC (Window B)",
                    level=2)
    cmd(doc, f"{CD}python3 scopus/at.py --point-here")
    expected(doc, "Six settings accepted, then read back, ending in:")
    code(doc, '  AT+SDVRSRVGET\n'
              '      +SDVRSRVGET:"192.168.2.3","",8080,"/upload",http\n'
              '      OK\n'
              '\n'
              'OK — notifications to 192.168.2.3:9999, photos to '
              'http://192.168.2.3:8080/upload')
    note(doc, "It judges on the read-back, not on the OKs: a setting that is "
              "accepted and stored nowhere is exactly what this step exists "
              "to catch.")

    doc.add_heading("Step 4 — Tell the camera what to do (Window B)", level=2)
    doc.add_paragraph("Five commands, in this order:")
    for c in ["detect stop", "detect debounce 3000", "detect profile 0x01 0x07",
              "notify enable 0x30", "detect start"]:
        cmd(doc, f'{CD}python3 scopus/cam.py "{c}"')
    expected(doc, "Each answers with what it set, and ok:")
    code(doc, "detect profile: det_msk=0x01 action_msk=0x07\n"
              "detect profile ok\n"
              "notify enable: 0x00000030\n"
              "notify enable ok\n"
              "detect: started\n"
              "detect start ok")
    doc.add_paragraph("What the three numbers mean — these are the ones "
                      "people get wrong:")
    table(doc, ["Setting", "Meaning"],
          [["detect profile 0x01", "what to look for: bit0 people, bit1 "
                                   "vehicles."],
           ["detect profile 0x07", "what to do about it: bit0 save to SD, "
                                   "bit1 report over the network, bit2 upload "
                                   "the photo. 0x07 is the full product; the "
                                   "default 0 detects and does nothing."],
           ["notify enable 0x30", "which events to report: 0x10 people, 0x20 "
                                  "vehicles."],
           ["detect debounce 3000", "how long a scene must hold before it "
                                    "counts, in ms. Without it a person on "
                                    "the edge of frame produces an event "
                                    "every few seconds."]],
          widths=[1.9, 4.4])

    doc.add_heading("Step 5 — Live picture (Window C, optional)", level=2)
    doc.add_paragraph("On the bench PC's own screen — this will not work over "
                      "SSH:")
    cmd(doc, "n6cam-view")
    doc.add_paragraph("A window opens showing what the camera sees. Useful "
                      "for Step 8; skip it if you are working remotely.")

    doc.add_heading("Step 6 — Prove an event reaches the server", level=2)
    doc.add_paragraph("In Window B:")
    cmd(doc, f'{CD}python3 scopus/cam.py "detect simulate 3" --wait 20')
    expected(doc, "In Window B, the camera raising the event:")
    code(doc, 'detect simulate: 3 object(s)\n'
              '+SDVRNTF: {"ser":4194336,"num":19,"rsn":16,"rsd":3,'
              '"tim":"20260817094834",…}\n'
              'detect simulate ok')
    expected(doc, "and within a few seconds, in Window A, the event arriving:")
    code(doc, "[09:48:33] NOTIFICATION   from 192.168.2.2  ser=4194336 num=19 "
              "rsn=16 rsd=3\n"
              "             valid JSON, 9 fields: {\"ser\":4194336,\"num\":19,"
              "\"rsn\":16,\"rsd\":3,…}\n"
              "[09:48:36] NOTIFICATION   from 192.168.2.2  ser=4194336 num=20 "
              "rsn=16 rsd=0")
    doc.add_paragraph("Two events, and both matter:")
    table(doc, ["Field", "Means"],
          [["rsd=3", "three people are there — the scene filled."],
           ["rsd=0", "they are gone — the scene emptied. Sent a few seconds "
                     "later, and its absence is a fault: a scene that never "
                     "reports empty leaves the operator watching a stale "
                     "count."],
           ["rsn=16", "the reason is People (0x10). 64 is a photo upload."],
           ["num", "the event's number. It goes up by one each time and never "
                   "repeats — the same number twice is a fault."]],
          widths=[1.1, 5.2])
    note(doc, "Window A is the pass condition, not Window B. The camera "
              "printing +SDVRNTF only proves it composed an event; the line "
              "in Window A proves it crossed both devices and arrived.")

    doc.add_heading("Step 7 — Prove a photo reaches the server", level=2)
    cmd(doc, f'{CD}python3 scopus/cam.py "photo upload" --wait 40')
    expected(doc, "In Window B:")
    code(doc, "photo upload: capturing -> SDVR+SENDBIN ref=6 "
              "name=4194336_17082026_094906.rdy\n"
              "photo upload ok\n"
              "+SDVRSRVR: OK\n"
              '+SDVRUPL: END,"upload_ref6"')
    expected(doc, "In Window A, about 15 seconds later:")
    code(doc, "[09:49:23] UPLOAD         from 192.168.2.2  118853 bytes  -> "
              "094923_photo\n"
              "[09:49:23]                JPEG, complete   path=/upload\n"
              "[09:49:23]                X-Filename: photo\n"
              "[09:49:23]                X-Filesize: 118853\n"
              "[09:49:23]                X-Ref: 6")
    note(doc, "JPEG, complete is the line to read. It means the file starts "
              "FFD8 and ends FFD9 — a transfer that was cut short still POSTs "
              "happily and would otherwise look like a pass.")
    doc.add_paragraph("Then look at the file itself, in Window B:")
    cmd(doc, "ls -lt ~/scopus-received | head -3")
    expected(doc)
    code(doc, "-rw-rw-r-- 1 user user 118853 Aug 17 09:49 094923_photo\n"
              "-rw-rw-r-- 1 user user    375 Aug 17 09:49 notifications.log")
    cmd(doc, "xdg-open ~/scopus-received/*_photo")
    doc.add_paragraph("The picture must open and show what the camera is "
                      "pointed at. A file of the right size that will not "
                      "open is a failure, not a pass.")
    note(doc, "About 100–130 KB is normal for 800×600. A few hundred bytes "
              "means the upload was cut short — report it with the ref "
              "number.")

    doc.add_heading("Step 8 — Real people", level=2)
    doc.add_paragraph(
        "Steps 6 and 7 used a simulated detection so the delivery path could "
        "be tested on its own. This step is the product: walk in front of "
        "the camera, stay about three seconds, then step out of frame.")
    expected(doc, "In Window A, without you typing anything — the same three "
                  "lines as Steps 6 and 7, with the counts the camera "
                  "actually saw:")
    code(doc, "[10:02:11] NOTIFICATION   from 192.168.2.2  ser=4194336 "
              "num=22 rsn=16 rsd=1\n"
              "[10:02:19] NOTIFICATION   from 192.168.2.2  ser=4194336 "
              "num=23 rsn=16 rsd=0\n"
              "[10:02:26] UPLOAD         from 192.168.2.2  117204 bytes  -> "
              "100226_photo")
    note(doc, "This is the one block in this document that is illustrative "
              "rather than captured: it needs a person in front of the "
              "camera. The line format is the real one from Steps 6 and 7; "
              "the numbers are what a single person walking past produces.")
    doc.add_paragraph("rsd is the number of people the camera counted. One "
                      "person must give rsd=1, two must give rsd=2.")
    note(doc, "Give it a few seconds either side — the debounce set in Step 4 "
              "is deliberately holding the scene for 3 seconds before it "
              "counts. Events that follow you within a second of moving mean "
              "the debounce did not take.")

    doc.add_heading("Step 9 — The box being moved", level=2)
    doc.add_paragraph(
        "Everything up to here is about what the camera sees. This step is "
        "about the unit itself being moved — picked up, knocked or tilted — "
        "which the board's own movement sensor reports, separately from "
        "anything in the picture.")
    doc.add_paragraph(
        "Step 4 enabled people and vehicles only (0x30). Turn the two motion "
        "events on as well, and shorten the wait before a stop is reported:")
    cmd(doc, f'{CD}python3 scopus/cam.py "notify enable 0x3f"')
    cmd(doc, f'{CD}python3 scopus/cam.py "motion sense 50 10"')
    doc.add_paragraph(
        "Then move the box: pick it up, put it down, or tap it firmly. If "
        "nobody is next to the unit, run this instead — the sensor pushes "
        "its own measuring element, which is a real movement:")
    cmd(doc, f'{CD}python3 scopus/cam.py "motion selftest"')
    expected(doc, "In Window A, without you typing anything further:")
    code(doc, "[17:02:17] NOTIFICATION   from 192.168.2.2  ser=4194336 "
              "num=0 rsn=2 rsd=688\n"
              "[17:02:33] NOTIFICATION   from 192.168.2.2  ser=4194336 "
              "num=1 rsn=4 rsd=15")
    table(doc, ["Field", "Meaning"],
          [["rsn=2", "Motion start — the unit began moving."],
           ["rsn=4", "Motion stop — it has been still for the timeout you "
                     "set (10 seconds above)."],
           ["rsd", "On start, how big the disturbance was, in mg. On stop, "
                   "how long the movement lasted, in seconds."],
           ["mtn", "1 while the box is moving. It appears on every "
                   "notification, so an event sent during a move says so."]])
    note(doc, "Walking in front of the camera must NOT produce rsn=2 or "
              "rsn=4. If it does, report it: motion means the box moving, "
              "not the scene changing.")
    doc.add_paragraph("Put the settings back when you are done:")
    cmd(doc, f'{CD}python3 scopus/cam.py "motion sense 50 30"')
    cmd(doc, f'{CD}python3 scopus/cam.py "notify enable 0x30"')

    doc.add_heading("Step 10 — Finish", level=2)
    doc.add_paragraph("Press Ctrl-C in Window A. Then, if you are done "
                      "testing, stop the detector:")
    cmd(doc, f'{CD}python3 scopus/cam.py "detect stop"')

    doc.add_heading("Part 1 — pass criteria", level=2)
    table(doc, ["#", "What must be true"],
          [["1", "preflight.py ends in READY."],
           ["2", "--point-here reads back this PC's address on both "
                 "endpoints."],
           ["3", "A simulated detection arrives in Window A as valid JSON "
                 "with rsd=3."],
           ["4", "The matching rsd=0 arrives a few seconds later."],
           ["5", "A photo arrives, is reported JPEG, complete, and opens."],
           ["6", "A real person in front of the camera produces the same two "
                 "events and a photo, with rsd matching the number of "
                 "people."],
           ["7", "Moving the box (or motion selftest) produces rsn=2 and "
                 "then rsn=4, and walking in front of the camera produces "
                 "neither."],
           ["8", "Every event has a different num. The same num twice is a "
                 "failure — report it."]],
          widths=[0.4, 5.9])

    # ══ PART 2 ═════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("3. Part 2 — The same over the mobile network, through "
                    "the customer's server", level=1)
    doc.add_paragraph(
        "Same camera, same modem, same events and the same photo — but the "
        "cable plays no part. The unit sends over the mobile network to the "
        "server, and the server sends commands back to the unit over an "
        "encrypted channel the unit holds open. Nothing is plugged in at "
        "either end.")
    note(doc, "The unit has no address anyone can dial — the operator's NAT "
              "sits in front of it — so it opens a connection outwards to the "
              "broker and holds it open, and commands are pushed down it. "
              "That is why there is a broker at all.")

    doc.add_heading("What is already set up on the server", level=2)
    doc.add_paragraph("Nothing in this part installs anything. It is all "
                      "running already:")
    table(doc, ["Piece", "Where", "What it is"],
          [["Receiver", f"{SRV}:{SRV_HTTP_PORT}",
            "POST /upload takes the photos, POST /notify takes the events. "
            f"Both land in {UPLOAD_DIR}."],
           ["Broker (mosquitto)", f"{SRV}:{MQTT_PORT}",
            "The command channel. TLS, and it refuses any client that does "
            "not present a certificate."],
           ["Certificates", CERTS,
            "ca.crt, client.crt, client.key. The same three are already on "
            "the modem under /data/sdvr/certs."],
           ["The unit's name", "its IMEI",
            "Every topic below is built from it. Step 2 reads it off the "
            "device — do not assume it."]],
          widths=[1.2, 1.6, 3.5])

    doc.add_heading("Step 1 — Point the unit at the server (once, on the "
                    "bench)", level=2)
    doc.add_paragraph(
        "This is the only step that needs the cable, and only because the "
        "unit has to be told where to send. A unit that has been configured "
        "once keeps these settings across reboots — skip to Step 2 if you "
        "are re-testing the same unit.")
    cmd(doc, f"{CD}python3 scopus/at.py --point-cloud {SRV} \\\n"
             f"  --http-port {SRV_HTTP_PORT} --path /upload "
             f"--notify-path /notify --apn {APN}")
    expected(doc, "The settings read back, then the link coming up:")
    code(doc, f'  AT+SDVRSRVGET\n'
              f'      +SDVRSRVGET:"{SRV}","",{SRV_HTTP_PORT},"/upload",http\n'
              f'      OK\n'
              f'\n'
              f'Waiting for the cellular link (up to 90 s)…\n'
              f'  AT+SDVRNET?\n'
              f'      +SDVRNET: 1,1,1,1,"Partner IL","LTE","internet",'
              f'"rmnet_data0",…\n'
              f'\n'
              f'OK — cellular link up. Notifications to '
              f'http://{SRV}:{SRV_HTTP_PORT}/notify,\n'
              f'photos to http://{SRV}:{SRV_HTTP_PORT}/upload')
    doc.add_paragraph("The four numbers after +SDVRNET: are the whole mobile "
                      "picture, and the fourth is the one that matters:")
    table(doc, ["Position", "Must be 1", "If it is 0"],
          [["1st  radio", "radio on", "the modem is not on the air."],
           ["2nd  registered", "registered on the network",
            "no coverage, or the SIM is in the holder the modem is not "
            "reading — check AT+SDVRSIM?."],
           ["3rd  session", "data session up",
            "wrong APN, or no data on the SIM."],
           ["4th  route", "there is a way out",
            "registered but with nowhere to send. Everything else looks "
            "healthy and nothing arrives — this is the flag to read first."]],
          widths=[1.3, 1.7, 3.3])

    doc.add_heading("Step 2 — Check the unit is on the command channel",
                    level=2)
    cmd(doc, f'{CD}python3 scopus/at.py "AT+SDVRMQTT?"')
    expected(doc)
    code(doc, f'+SDVRMQTT: 1,1,"{SRV}",{MQTT_PORT},"{IMEI}",6,7,1')
    doc.add_paragraph(
        "The first number is 1 when the channel is switched on, the second "
        "is 1 when it is actually connected. The text in the middle is the "
        "unit's name — write it down, every command below needs it.")
    note(doc, "If the second number is 0, give it half a minute and look "
              "again: it retries on its own, 5 s then 10, 20 and so on up to "
              "a minute. If it stays 0, the mobile side is the suspect, not "
              "this channel — go back to Step 1 and read the four flags.")
    doc.add_paragraph("If the channel is switched off, switch it on:")
    cmd(doc, f'{CD}python3 scopus/at.py '
             f'"AT+SDVRMQTTSRV=\\"{SRV}\\",{MQTT_PORT}"')
    cmd(doc, f'{CD}python3 scopus/at.py "AT+SDVRMQTT=1"')

    doc.add_heading("Step 3 — Watch the unit (Window A, on the server)",
                    level=2)
    doc.add_paragraph("Leave this running. It shows every command that goes "
                      "out and every answer that comes back. Replace <IMEI> "
                      "with the name from Step 2.")
    cmd(doc, f"mosquitto_sub -h {SRV} -p {MQTT_PORT} \\\n"
             f"  {MQTT_AUTH} \\\n"
             f"  -t 'scopus/<IMEI>/#' -v")
    expected(doc, "Straight away, one line:")
    code(doc, f"scopus/{IMEI}/status online")
    note(doc, "That line is stored by the broker, so it appears the moment "
              "you subscribe rather than when the unit next speaks. If it "
              "says offline, the unit's connection has dropped and the broker "
              "is saying so on its behalf.")

    doc.add_heading("Step 4 — Send a command (Window B, on the server)",
                    level=2)
    cmd(doc, f"mosquitto_pub -h {SRV} -p {MQTT_PORT} \\\n"
             f"  {MQTT_AUTH} \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'version'")
    expected(doc, "In Window A, within a few seconds, two lines:")
    code(doc, f"scopus/{IMEI}/cmd version\n"
              f"scopus/{IMEI}/rsp Application: 01.08.2593089169 "
              f"Build: Aug 16 2026 16:24:44 version ok")
    doc.add_paragraph("The first is your own command coming back past you; "
                      "the second is the unit's answer. Anything you can type "
                      "at the camera console works here.")
    doc.add_paragraph("A command that starts with AT is answered by the "
                      "modem instead of the camera. This is how a unit in the "
                      "field gets diagnosed without anyone travelling to it:")
    cmd(doc, f"mosquitto_pub -h {SRV} -p {MQTT_PORT} \\\n"
             f"  {MQTT_AUTH} \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'AT+CSQ'")
    expected(doc)
    code(doc, f"scopus/{IMEI}/rsp AT+CSQ\n"
              f"+CSQ: 29,99\n"
              f"OK")
    table(doc, ["Command", "What it tells you"],
          [["AT+CSQ", "signal strength. First number: under 10 poor, over 20 "
                      "good, 99 no signal."],
           ["AT+SDVRNET?", "the whole mobile picture — the four flags from "
                           "Step 1, operator, APN, address."],
           ["AT+SDVRMQTT?", "this channel's own state and counters."],
           ["AT+SDVRSIM?", "which SIM holder the unit is reading."],
           ["mdm stats", "the internal cable between the two boards — read it "
                         "when events stop arriving."]],
          widths=[1.6, 4.7])

    doc.add_heading("Step 5 — The whole product, from the server", level=2)
    doc.add_paragraph("This is the test that matters: the command goes out "
                      "over the mobile network, and the result comes back "
                      "over it. First an event:")
    cmd(doc, f"mosquitto_pub -h {SRV} -p {MQTT_PORT} \\\n"
             f"  {MQTT_AUTH} \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'detect simulate 3'")
    expected(doc, "In Window A:")
    code(doc, f"scopus/{IMEI}/cmd detect simulate 3\n"
              f"scopus/{IMEI}/rsp detect simulate: 3 object(s) "
              f"detect simulate ok")
    doc.add_paragraph("Then a photo:")
    cmd(doc, f"mosquitto_pub -h {SRV} -p {MQTT_PORT} \\\n"
             f"  {MQTT_AUTH} \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'photo upload'")
    expected(doc, "In Window A:")
    code(doc, f"scopus/{IMEI}/rsp photo upload: capturing -> SDVR+SENDBIN "
              f"ref=7 name=4194336_17082026_095119.rdy photo upload ok")

    doc.add_heading("Step 6 — Check what actually landed on the server",
                    level=2)
    doc.add_paragraph("The answers in Window A only say the unit accepted the "
                      "command. This step is the pass condition. In Window B:")
    cmd(doc, f"tail -4 {UPLOAD_DIR}/_uploads.log")
    expected(doc, "The event and the photo, arriving from the internet:")
    code(doc, "[2026-08-17T09:50:42] 213.8.185.178:36680 POST /notify "
              "name=upload.bin bytes=101\n"
              "[2026-08-17T09:50:46] 213.8.185.178:36681 POST /notify "
              "name=upload.bin bytes=101\n"
              "[2026-08-17T09:51:33] 213.8.185.178:44805 POST /upload "
              "name=photo bytes=118768\n"
              "[2026-08-17T09:51:33] 213.8.185.178:44812 POST /notify "
              "name=upload.bin bytes=101")
    note(doc, "213.8.185.178 is the mobile operator's address, not the "
              "unit's and not this PC's. That is the proof it came in over "
              "the network rather than the cable.")
    doc.add_paragraph("Now the event body and the photo themselves:")
    cmd(doc, f"cat {UPLOAD_DIR}/upload.bin")
    expected(doc)
    code(doc, '{"ser":4194336,"num":23,"rsn":16,"rsd":0,'
              '"tim":"20260817095044","mtn":0,"mod":"","bat":0.0,"vol":0.0}')
    note(doc, "rsd=0 here is correct, not a miscount: the file holds the "
              "LAST event, and the last of the pair is the scene emptying. "
              "To see the rsd=3 that went with it, run this command in the "
              "four seconds between the two log lines, or read the pair off "
              "Window A in Part 1 where nothing is overwritten.")
    cmd(doc, f"ls -l {UPLOAD_DIR}/photo && xdg-open {UPLOAD_DIR}/photo")
    expected(doc)
    code(doc, f"-rw-r--r-- 1 user user 118768 Aug 17 09:51 "
              f"{UPLOAD_DIR}/photo")
    doc.add_paragraph("The picture must open and show what the camera is "
                      "pointed at.")
    note(doc, "Both files are overwritten by each new arrival — the receiver "
              "keeps one photo and one event body, and the history is in "
              "_uploads.log. Check the timestamp, not just that a file is "
              "there.", warn=True)

    doc.add_heading("Part 2 — pass criteria", level=2)
    table(doc, ["#", "What must be true"],
          [["1", "AT+SDVRNET? shows 1,1,1,1 — in particular the fourth flag."],
           ["2", "AT+SDVRMQTT? shows 1,1 — switched on and connected."],
           ["3", "The status topic says online as soon as you subscribe."],
           ["4", "A camera command (version) is answered."],
           ["5", "A modem command (AT+CSQ) is answered."],
           ["6", "detect simulate 3 puts two POST /notify lines in the "
                 "server's log, a few seconds apart, and the body is valid "
                 "JSON."],
           ["7", "photo upload sent from the server puts a JPEG on the "
                 "server that opens."],
           ["8", "Every command is answered exactly ONCE. Two identical "
                 "answers to one command is a failure — report it."]],
          widths=[0.4, 5.9])

    # ══ Troubleshooting ════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("4. If something fails", level=1)
    doc.add_paragraph("In order — the first row that matches is usually the "
                      "answer.")
    table(doc, ["What you see", "What it means and what to do"],
          [["preflight says the camera port is missing, or cam.py answers "
            "nothing",
            "The camera's USB connection has stopped draining, which is not "
            "the same as the camera being down. Cut the port's power and "
            "wait: sudo uhubctl -l 3-7 -p 1 -a cycle --delay 3 && sleep 15. "
            "This does not reboot the camera."],
           ["cam.py works but no event reaches the server",
            "The internal cable between the two boards is the suspect. Run "
            "cam.py \"mdm AT\" — it must answer OK. If it does not, run "
            "cam.py \"mdm relink\" and try again."],
           ["Events arrive but no photo",
            "Check the profile from Step 4: uploading the photo is bit2, so "
            "the action mask has to be 0x07 and not 0x03."],
           ["Nothing arrives in Part 2, and Part 1 was fine",
            "The product is fine and the network is not. Read the four flags "
            "of AT+SDVRNET? — the fourth being 0 is the usual answer, and "
            "means registered with nowhere to send."],
           ["AT+SDVRMQTT? shows 1,0 and stays there",
            "Switched on but not connecting. Almost always no mobile data "
            "rather than anything to do with this channel — check Step 1. If "
            "the mobile side is healthy, the certificates are next: the "
            "modem needs all three files under /data/sdvr/certs."],
           ["mosquitto_sub exits, or says connection refused",
            "The broker would not accept your certificate. Check you passed "
            "all three of --cafile, --cert and --key — leaving out --cert is "
            "the usual mistake, and the broker refuses without explaining."],
           ["Your command appears on .../cmd but no answer comes back",
            "It reached the unit and the answer did not return. If it was a "
            "camera command, the internal cable is the suspect — run "
            "mdm AT. If it was an AT command, report it."],
           ["The answer arrives twice",
            "A fault, not a quirk. Report it with the command you sent."],
           ["The answer stops mid-sentence with [truncated]",
            "Expected: a very long answer is cut after about 1.8 KB. Ask for "
            "less at a time."],
           ["The modem answers but AT+SDVRVER says an old version",
            "The app rolled back — a Legato install that was never marked "
            "good reverts on the next reboot. Report the version you see; "
            "reinstalling is a developer step, not a QA one."]],
          widths=[2.2, 4.1])

    doc.add_heading("Known behaviour — not your fault", level=2)
    for a in [
        "The receiver keeps one file called photo and one called upload.bin, "
        "overwriting both on every arrival. Go by the timestamp and by "
        "_uploads.log.",
        "An event body's mod, bat and vol fields are empty or zero. Nothing "
        "populates them yet; the transport carries them at full size.",
        "AT+SDVRNTFPROTO=2 (events over the command channel instead of a "
        "POST) answers +SDVRERR: 13 on firmware 1.12.0. Use the POST path in "
        "Step 1 — it is what this document tests and it works.",
        "The first command after a long idle gap is sometimes swallowed and "
        "the tools retry it. A single repeat is normal; a command that never "
        "answers is not.",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    # ══ Cheat sheet ════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("5. Cheat sheet — the whole flow, in order", level=1)
    doc.add_paragraph("Part 1, on the bench PC:").bold = True
    code(doc,
         f"{CD}python3 scopus/preflight.py\n"
         f"{CD}python3 scopus/test_server.py --http-port 8080 --udp-port 9999 "
         f"\\\n"
         f"     --dir ~/scopus-received --from-modem 192.168.2.2 --fresh   "
         f"# Window A\n"
         f"{CD}python3 scopus/at.py --point-here\n"
         f'{CD}python3 scopus/cam.py "detect stop"\n'
         f'{CD}python3 scopus/cam.py "detect debounce 3000"\n'
         f'{CD}python3 scopus/cam.py "detect profile 0x01 0x07"\n'
         f'{CD}python3 scopus/cam.py "notify enable 0x30"\n'
         f'{CD}python3 scopus/cam.py "detect start"\n'
         f'{CD}python3 scopus/cam.py "detect simulate 3" --wait 20\n'
         f'{CD}python3 scopus/cam.py "photo upload" --wait 40\n'
         f'{CD}python3 scopus/cam.py "notify enable 0x3f"\n'
         f'{CD}python3 scopus/cam.py "motion sense 50 10"\n'
         f'{CD}python3 scopus/cam.py "motion selftest"      '
         f'# or move the box by hand\n'
         f'{CD}python3 scopus/cam.py "motion sense 50 30"\n'
         f"ls -lt ~/scopus-received | head -3")
    doc.add_paragraph()
    doc.add_paragraph("Part 2 — device side, once, on the bench PC:").bold = True
    code(doc,
         f"{CD}python3 scopus/at.py --point-cloud {SRV} --http-port "
         f"{SRV_HTTP_PORT} \\\n"
         f"     --path /upload --notify-path /notify --apn {APN}\n"
         f'{CD}python3 scopus/at.py "AT+SDVRMQTT?"        '
         f'# note the IMEI')
    doc.add_paragraph()
    doc.add_paragraph("Part 2 — server side (ID is the IMEI):").bold = True
    code(doc,
         f"C={CERTS}\n"
         f"ID={IMEI}\n"
         f"\n"
         f"mosquitto_sub -h {SRV} -p {MQTT_PORT} --cafile $C/ca.crt \\\n"
         f"     --cert $C/client.crt --key $C/client.key \\\n"
         f"     -t \"scopus/$ID/#\" -v                    # Window A\n"
         f"\n"
         f"mosquitto_pub -h {SRV} -p {MQTT_PORT} --cafile $C/ca.crt \\\n"
         f"     --cert $C/client.crt --key $C/client.key \\\n"
         f"     -t \"scopus/$ID/cmd\" -q 1 -m 'version'\n"
         f"\n"
         f"# then, one at a time, in place of 'version':\n"
         f"#   AT+CSQ            AT+SDVRNET?        mdm stats\n"
         f"#   detect simulate 3                    photo upload\n"
         f"\n"
         f"tail -4 {UPLOAD_DIR}/_uploads.log\n"
         f"cat {UPLOAD_DIR}/upload.bin\n"
         f"ls -l {UPLOAD_DIR}/photo")
    doc.add_paragraph()
    doc.add_paragraph("The four topics, if you want to look at them "
                      "directly:")
    table(doc, ["Topic", "Direction", "Carries"],
          [["scopus/<IMEI>/cmd", "to the unit", "one command per message, "
                                                "plain text"],
           ["scopus/<IMEI>/rsp", "from the unit", "that command's output"],
           ["scopus/<IMEI>/ntf", "from the unit", "events, when the unit is "
                                                  "set to report over this "
                                                  "channel"],
           ["scopus/<IMEI>/status", "from the unit",
            "online / offline — the broker publishes offline itself if the "
            "unit drops"]],
          widths=[1.8, 1.2, 3.3])

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
