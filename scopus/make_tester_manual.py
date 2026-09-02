#!/usr/bin/env python3
"""Generate Scopus_Tester_Manual.docx from the procedure below.

The manual is generated rather than hand-edited so the commands in it stay in
one place with the rest of the suite, and so a change to the procedure is a
reviewable diff instead of a binary blob nobody can read. Re-run after editing:

    python3 scopus/make_tester_manual.py

It is written for a tester who is testing the product, not administering Linux:
every command is one complete line that can be copied as it stands, and every
"Expected" block is what the bench actually printed when this procedure was
walked line by line — not what it ought to print.

The commands all start with `cd ~/work/itpnovex/edgeai &&` on purpose. It is
redundant once you are in that directory, and it means a tester who opened a
new window, or lost their place, can start from any line in the document and
have it work.
"""
import datetime
import os
import pathlib
import sys

# ScopusQA #11 — site data (bench host, receiver address) is read from
# scopus/bench.ini, which is untracked; bench.ini.template is the committed
# copy and carries placeholders only. See scopus/lib/settings.py.
import os as _os, sys as _sys
_sys.path.insert(0, _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "lib"))
from settings import S as _S

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

OUT = pathlib.Path(__file__).resolve().parent / "Scopus_Tester_Manual.docx"

MONO = "Consolas"
GREY = RGBColor(0x44, 0x44, 0x44)
RED = RGBColor(0xB0, 0x00, 0x00)
GREEN = RGBColor(0x1E, 0x6B, 0x2E)
BLUE = RGBColor(0x1F, 0x3D, 0x7A)

REPO = "~/work/itpnovex/edgeai"
CD = f"cd {REPO} && "

# The machine this document is written for: the one the camera and the modem
# are physically plugged into. Naming it matters — "the bench PC" is ambiguous
# the moment there is more than one, and a tester who runs these commands on
# the wrong machine gets "no such device" for reasons no troubleshooting table
# can explain.
BENCH_HOST = _S.get("bench", "host") or "<your bench PC>"
BENCH_ADDR = "100.115.215.6"
BENCH_SSH_PORT = 4322
BENCH_USER = "user"

# The APN this bench's SIM uses. A placeholder here would be a step the tester
# cannot complete, and the wrong APN is indistinguishable from a dead SIM.
BENCH_APN = "internet"

# Remote command channel (section 18). The broker runs on this same PC and is
# reached on its public address, because that is the address the unit dials
# from the mobile network — the LAN address is not routable from there.
MQTT_HOST = _S.require("server", "host")
MQTT_PORT = 5912
MQTT_CERT_DIR = "/opt/sdvr-server/certs"


def code(doc, text, size=9):
    """A command / output block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(size)
    return p


def cmd(doc, text):
    """A line the tester types. Bold so it stands out from expected output."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = BLUE
    return p


def note(doc, text, warn=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RED if warn else GREY
    return p


def expected(doc, text="Expected:"):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(2)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    if widths:
        # Three things have to agree or the widths are ignored: autofit off,
        # every cell's own width, and w:tblGrid — which is the one that
        # actually governs under a fixed layout, and which python-docx leaves
        # at equal columns however many cell widths you set.
        t.autofit = False
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
        grid = t._tbl.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            "tblGrid")
        if grid is not None:
            for col, w in zip(grid, widths):
                col.set("{http://schemas.openxmlformats.org/wordprocessingml/"
                        "2006/main}w", str(int(w * 1440)))
    return t


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    doc.add_heading("Scopus — Tester Manual", 0)
    p = doc.add_paragraph()
    p.add_run("End-to-end acceptance test: the camera sees people → the modem "
              "sends the event → this PC receives it.").bold = True
    doc.add_paragraph(
        f"Generated {datetime.date.today().isoformat()} from "
        f"scopus/make_tester_manual.py. Every command below was run, in this "
        f"order, on the bench PC this document is written for, and each "
        f"“Expected” block is the output that came back. The one "
        f"exception is Step 9, which needs a person standing in front of the "
        f"camera; its example is a real live detection recorded on the same "
        f"bench.")

    # ── How to use this document ───────────────────────────────────────
    doc.add_heading("1. How to use this document", level=1)
    doc.add_paragraph(
        "Work through it from top to bottom. Each step gives you one command "
        "to type and shows what should come back. If what you see does not "
        "match, stop and look up the symptom in section 14 — carrying on "
        "after a failed step produces confusing results later.")
    for a in [
        "Lines in bold blue are what you type. Copy the whole line, including "
        "the part before &&.",
        "Everything else in the boxes is what the computer prints back. It "
        "will not match character for character — times, counters and the "
        "picture change — but the shape and the key values must match.",
        "You need two Terminal windows. Open one with Ctrl+Alt+T, and a "
        "second the same way. This document calls them Window A and Window B.",
        "Nothing here needs a password, and nothing here changes the "
        "product's software.",
    ]:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph("What the two windows are for:")
    table(doc, ["Window", "What it does", "When"],
          [["A", "Runs the receiving server — this is 'the customer's server' "
                 "that the product sends events and photos to. It prints "
                 "everything that arrives.",
            "Started in Step 3, then left alone until the end"],
           ["B", "Everything else: the checks and the commands to the camera "
                 "and the modem.",
            "Used for every other step"]])

    doc.add_paragraph()
    doc.add_paragraph(
        "Roughly 20 minutes end to end, most of it waiting for the photo to "
        "cross the internal cable between the two boards.")

    doc.add_paragraph()
    doc.add_paragraph("There are two ways to run this test:").bold = True
    table(doc, ["Mode", "How the product sends", "Where to look"],
          [["Cable", "Over the USB cable to this PC, which plays the part of "
                     "the customer's server. Nothing leaves the room.",
            "Steps 1 to 12, in order — this is the default"],
           ["Cellular", "Over the mobile network to the customer's server, "
                        "which this PC then reads. This is what a deployed "
                        "unit does.",
            "Scopus_QA_Flow.docx, Part 2 — a document of its own"]])
    doc.add_paragraph(
        "Run the cable mode first, even if cellular is what you were asked "
        "to test. It uses the same camera, the same modem and the same "
        "internal cable, so if it fails, cellular was never going to work "
        "and you have already found out where — which is the one thing the "
        "cellular mode cannot tell you, because a fault in the product and a "
        "fault in the mobile network look identical from the far end.")

    # ── What this proves ───────────────────────────────────────────────
    doc.add_heading("2. What this test proves", level=1)
    doc.add_paragraph(
        "Scopus is two devices in one box: an N6Cam (camera + neural network) "
        "and a WP76 modem, joined by an internal serial cable. Each half can "
        "look perfectly healthy while the product does nothing, because "
        "nothing joins them. This procedure follows real events all the way "
        "out to this PC:")
    code(doc,
         "  people in front of the camera\n"
         "      -> camera's neural network detects them\n"
         "          -> camera tells the modem over the internal cable\n"
         "              -> modem sends a UDP message to this PC\n"
         "\n"
         "  photo capture\n"
         "      -> camera sends the JPEG to the modem\n"
         "          -> modem uploads it to this PC over HTTP")
    doc.add_paragraph(
        "The test passes on what arrives on this PC — a JSON event and a JPEG "
        "file you can open. A log line on the device saying it sent them is "
        "not a pass; that only proves the device spoke, not that anything "
        "heard it.")
    doc.add_paragraph("The steps prove three separate things, which is why "
                      "there are three of them:")
    table(doc, ["Step", "Proves"],
          [["7 — known picture",
            "the neural network counts people correctly (the picture has a "
            "known number in it, so the answer is checkable)"],
           ["8 — test event",
            "an event gets from the camera, through the modem, onto this PC, "
            "and arrives as valid data"],
           ["9 — real people",
            "the two above working together on a real detection: the product "
            "doing its actual job"]])

    # ── Equipment ──────────────────────────────────────────────────────
    doc.add_heading("3. What is already set up", level=1)
    doc.add_paragraph(
        "This is all in place on the bench PC. You do not have to do "
        "anything with it; it is here so you can recognise the parts by name "
        "when a step mentions them.")
    table(doc, ["Item", "Detail"],
          [["This PC", f"{BENCH_HOST} — the PC the two boards are plugged "
                       f"into, and the one every command in this document is "
                       f"typed on. It plays the part of the customer's "
                       f"server. Its address on the modem's network is "
                       f"192.168.2.3"],
           ["N6Cam", "The camera board. Connected to this PC by USB"],
           ["WP76 modem", "The modem board. Connected to this PC by USB (for "
                          "commands) and Ethernet (192.168.2.2, for the data "
                          "it sends)"],
           ["Internal cable", "Camera to modem. This is the link the product "
                              "depends on and the one this test exercises"],
           ["Software", f"Already installed at {REPO}"]])
    doc.add_paragraph()
    doc.add_paragraph(
        "Every command in this document is typed on that PC. Sit at it if "
        "you can. If you are working from your own machine instead, open a "
        "terminal there and connect first — everything afterwards is "
        "identical:")
    cmd(doc, f"ssh -p {BENCH_SSH_PORT} {BENCH_USER}@{BENCH_ADDR}")
    note(doc, f"The login password is the same as the machine's name, "
              f"{BENCH_HOST}. If the connection times out, you are not on a "
              f"network that can see this PC — that is a network problem to "
              f"raise, not a test failure.")

    # ── Step 1 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 1 — Close anything else using the devices",
                    level=1)
    doc.add_paragraph(
        "If another program has the camera's or the modem's serial port open, "
        "it takes the replies this test is waiting for, and every step below "
        "fails with “no answer” while the hardware is perfectly "
        "healthy. This has cost real time on this bench.")
    doc.add_paragraph("Close, if any of them are open:")
    for a in ["The Serial Monitor panel in VS Code (this is the usual "
              "culprit — closing the panel is enough, but closing VS Code is "
              "surer)",
              "Any terminal window running picocom, minicom, tio or screen",
              "A test server left running from an earlier test"]:
        doc.add_paragraph(a, style="List Bullet")
    note(doc, "Step 2 checks this for you and names the program if it finds "
              "one, so you do not have to guess.")

    # ── Step 2 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 2 — Check the bench is ready", level=1)
    doc.add_paragraph("In Window B:")
    cmd(doc, f"{CD}python3 scopus/preflight.py")
    expected(doc)
    code(doc,
         "Scopus bench pre-flight\n\n"
         "  [PASS] N6Cam shell port  —  /dev/ttyACM2\n"
         "  [PASS] Modem AT port  —  /dev/ttyUSB0\n"
         "  [PASS] N6Cam port is free  —  nobody else has it open\n"
         "  [PASS] Modem port is free  —  nobody else has it open\n"
         "  [PASS] This PC is on the modem's network  —  this PC is 192.168.2.3\n"
         "  [PASS] Modem answers on the network  —  192.168.2.2\n"
         "  [PASS] Server ports are free  —  TCP 8080, UDP 9999\n"
         "  [PASS] Modem SDVR app answers  —  version 1.7.1\n"
         "  [PASS] Modem firmware is new enough  —  1.7.1 >= 1.7.0\n"
         "  [PASS] Camera shell answers  —  uptime replied\n"
         "  [PASS] Camera can reach the modem  —  internal cable carries commands\n"
         "  [PASS] Test image present  —  /home/user/work/itpnovex/edgeai/images/3_people.jpg\n\n"
         "READY — the bench is fit to test on. Go to Step 3 of the manual.")
    note(doc, "Every line must say PASS. If one says FAIL, the yellow text "
              "underneath it says what to do; do that and run the command "
              "again. Do not start the test with a FAIL showing — the step "
              "that fails later will look like a product fault and will not "
              "be one.", warn=True)
    note(doc, "'/dev/ttyACM2' may read ttyACM1 on your run. That is normal — "
              "the camera moves between the two — and nothing in this "
              "document depends on which one it is.")

    # ── Step 3 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 3 — Start the receiving server (Window A)", level=1)
    doc.add_paragraph(
        "This is the server the product uploads to. It listens for both "
        "things the device sends — events (UDP) and photos (HTTP) — and "
        "saves everything it receives so you can look at it afterwards.")
    doc.add_paragraph("In Window A:")
    cmd(doc, f"{CD}python3 scopus/test_server.py --http-port 8080 "
             f"--udp-port 9999 --dir ~/scopus-received --from-modem "
             f"192.168.2.2 --fresh")
    expected(doc)
    code(doc,
         "Scopus test server\n"
         "  receiving into /home/user/scopus-received\n"
         "  (an earlier run's files were moved to "
         "/home/user/scopus-received-old-20260806-151456)\n"
         "[15:14:56] listening      UDP  0.0.0.0:9999  (notifications)\n"
         "[15:14:56] listening      HTTP 0.0.0.0:8080 (photo uploads)\n"
         "  Ctrl-C to stop. Point the modem here with:\n"
         '    AT+SDVRNTFHOST="<this PC ip>"   AT+SDVRNTFPORT=9999\n'
         '    AT+SDVRHOSTIP="<this PC ip>"    AT+SDVRPORT=8080   '
         'AT+SDVRSRVRPATH="/upload"')
    doc.add_paragraph(
        "Leave this window running and visible for the rest of the test — "
        "several steps are judged by what appears here. Do not type in it. "
        "You will stop it with Ctrl-C at the very end.")
    doc.add_paragraph("What the options mean:")
    table(doc, ["Option", "Meaning"],
          [["--http-port 8080", "port the photos are uploaded to"],
           ["--udp-port 9999", "port the events arrive on"],
           ["--dir ~/scopus-received", "folder everything received is saved "
                                       "in"],
           ["--from-modem 192.168.2.2", "ignore anything that did not come "
                                        "from the modem. Unrelated traffic on "
                                        "this port has been mistaken for a "
                                        "passing test before"],
           ["--fresh", "start with an empty folder, so what is in it at the "
                       "end is this test's evidence and nothing else. An "
                       "earlier run's files are moved aside, not deleted"]])

    # ── Step 4 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 4 — Point the modem at this PC (Window B)", level=1)
    doc.add_paragraph(
        "The modem has to be told where to send things. This one command "
        "works out this PC's address on the modem's network, sets the five "
        "settings that need to agree with each other, and reads them back.")
    cmd(doc, f"{CD}python3 scopus/at.py --point-here")
    expected(doc)
    code(doc,
         "[modem /dev/ttyUSB0]\n"
         "Pointing the modem at this PC (192.168.2.3)\n\n"
         '  AT+SDVRNTFHOST="192.168.2.3"\n'
         "      OK\n"
         "  AT+SDVRNTFPORT=9999\n"
         "      OK\n"
         '  AT+SDVRHOSTIP="192.168.2.3"\n'
         "      OK\n"
         "  AT+SDVRPORT=8080\n"
         "      OK\n"
         '  AT+SDVRSRVRPATH="/upload"\n'
         "      OK\n\n"
         "Read back what the modem now has:\n"
         "  AT+SDVRNTFHOST?\n"
         '      +SDVRNTFHOST:"192.168.2.3"\n'
         "      OK\n"
         "  AT+SDVRNTFPORT?\n"
         "      +SDVRNTFPORT:9999\n"
         "      OK\n"
         "  AT+SDVRSRVGET\n"
         '      +SDVRSRVGET:"192.168.2.3","",8080,"/upload",http\n'
         "      OK\n\n"
         "OK — notifications to 192.168.2.3:9999, photos to "
         "http://192.168.2.3:8080/upload")
    doc.add_paragraph("The five settings, in case you need to recognise them:")
    table(doc, ["Setting", "Meaning"],
          [['AT+SDVRNTFHOST="192.168.2.3"', "where events go (this PC)"],
           ["AT+SDVRNTFPORT=9999", "port events go to"],
           ['AT+SDVRHOSTIP="192.168.2.3"', "where photos go (this PC)"],
           ["AT+SDVRPORT=8080", "port photos go to"],
           ['AT+SDVRSRVRPATH="/upload"', "the address on the server the photo "
                                         "is posted to"]])
    doc.add_paragraph()
    note(doc, "The last line must say OK. The command judges itself on the "
              "read-back rather than on the five OKs, because a setting that "
              "answers OK and stores nothing is exactly the fault this step "
              "exists to catch.")

    # ── Step 5 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 5 — Tell the camera to look for people", level=1)
    doc.add_paragraph("Two commands, in Window B:")
    cmd(doc, f'{CD}python3 scopus/cam.py "detect profile 0x01 0x03"')
    expected(doc)
    code(doc, "[camera /dev/ttyACM2] > detect profile 0x01 0x03\n"
              "detect profile: det_msk=0x01 action_msk=0x03\n"
              "detect profile ok")
    doc.add_paragraph()
    cmd(doc, f'{CD}python3 scopus/cam.py "notify enable 0xff"')
    expected(doc)
    code(doc, "[camera /dev/ttyACM2] > notify enable 0xff\n"
              "notify enable: 0x000000ff\n"
              "notify enable ok")
    doc.add_paragraph()
    table(doc, ["Command", "Meaning"],
          [["detect profile 0x01 0x03",
            "look for people (0x01); when you find some, both save a picture "
            "to the SD card and report it (0x03). The 'report' half is what "
            "makes an event reach this PC — with it off the camera detects "
            "silently"],
           ["notify enable 0xff", "allow every kind of event to be reported"]])
    doc.add_paragraph()
    note(doc, "These two settings go back to their defaults if the camera "
              "restarts. If a later step stops producing events, come back "
              "and run these two again — see section 15.")

    # ── Step 6 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 6 — Check the two boards are talking", level=1)
    doc.add_paragraph(
        "This sends a command from the camera, across the internal cable, to "
        "the modem, and brings the modem's answer back. Everything after this "
        "depends on that cable.")
    cmd(doc, f'{CD}python3 scopus/cam.py "mdm AT"')
    expected(doc)
    code(doc, "[camera /dev/ttyACM2] > mdm AT\n"
              "OK\n"
              "mdm AT ok")
    note(doc, "If this does not answer OK, stop here. The camera cannot reach "
              "the modem, and no step after this one can pass. Report it — "
              "and see section 14, which has the one thing worth trying "
              "first.", warn=True)

    # ── Step 7 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 7 — Check the neural network counts correctly",
                    level=1)
    doc.add_paragraph(
        "Rather than pointing the lens at real people — where nobody knows "
        "what the right answer was — this pushes a picture with a known "
        "number of people in it into the camera's neural network, and checks "
        "the number that comes back. It then puts the camera back on its own "
        "lens, which is what Step 9 needs. It takes about a minute.")
    cmd(doc, f"{CD}python3 scopus/inference_test.py")
    expected(doc)
    code(doc,
         "Camera on /dev/ttyACM2 — expecting 3 people in 3_people.jpg\n\n"
         "Attempt 1 of 4:\n"
         "  [1/5] stopping live detection\n"
         "  [2/5] injecting images/3_people.jpg\n"
         "        Camera port: /dev/ttyACM2\n"
         "        Frame: 256x256 RGB888 (196608 bytes, CRC32 0x13a92fb1)\n"
         "        frame upload: ok (196608 bytes, CRC 0x13a92fb1) frame upload ok  >\n"
         "        Uploaded. Next:  > frame run    (over the same CDC port)\n"
         "  [3/5] starting detection\n"
         "  [4/5] running inference on it\n"
         "        frame run: 3 detection(s), NN 87.9ms\n"
         "        [0] class=0 conf=0.78 bbox=(0.73,0.71,0.18,0.44)\n"
         "        [1] class=0 conf=0.71 bbox=(0.51,0.64,0.15,0.60)\n"
         "        [2] class=0 conf=0.84 bbox=(0.36,0.73,0.14,0.41)\n"
         "        frame run ok\n"
         "  [5/5] returning the camera to the live lens\n"
         "        frame: cleared (NN back to live camera)\n"
         "        frame clear ok\n\n"
         "PASSED — 3 people detected in 87.9 ms, which matches the picture. "
         "The camera is\nback on its live lens.")
    doc.add_paragraph("What to check:")
    for a in ["The last line says PASSED.",
              "Three detections, one per person in the picture.",
              "class=0 on every line — class 0 is 'person'. A different "
              "class number means it found something, but not a person.",
              "NN 87.9ms — how long the network took. Anything in the "
              "80–100 ms range is normal.",
              "Step [5/5] is there and says 'frame: cleared (NN back to live "
              "camera)'. Do not skip this line: it is the one that gives the "
              "camera its own lens back."]:
        doc.add_paragraph(a, style="List Bullet")
    note(doc, "Why that last line matters. While the picture is loaded the "
              "camera looks at it instead of through its lens, so Step 9 "
              "cannot pass, and the live view draws the picture's people on "
              "top of the real video, which looks like the detection has "
              "gone wrong. The step above clears it for you. If for any "
              "reason it did not say 'cleared', do it by hand before going "
              "on:", warn=True)
    cmd(doc, f'{CD}python3 scopus/cam.py "frame clear"')
    expected(doc)
    code(doc, "[camera /dev/ttyACM2] > frame clear\n"
              "frame: cleared (NN back to live camera)\n"
              "frame clear ok")
    note(doc, "The camera guards this itself as well, from the firmware of 9 "
              "August 2026 onwards. A loaded picture lapses on its own about "
              "two minutes after the last time it was used, and while it is "
              "loaded the live view says TEST PICTURE - NOT THE LENS with the "
              "seconds remaining, and the camera sends no detection events — "
              "a photograph is not allowed to be reported as people in the "
              "room. So the worst a forgotten picture can now cost you is a "
              "two-minute wait. Clear it anyway; knowing why the safety net "
              "is there is not a reason to land in it.")
    note(doc, "It should say Attempt 1 of 4 and stop there. A second "
              "attempt is not a failure in itself — only the final PASSED / "
              "FAILED line decides the step — but it used to be caused by a "
              "firmware fault fixed on 6 August 2026, so mention it in your "
              "report if you see one.")
    doc.add_paragraph(
        "There are other pictures to try, if you want more than one data "
        "point:")
    cmd(doc, f"{CD}python3 scopus/inference_test.py --image "
             f"images/2_people.jpg --expect 2")
    doc.add_paragraph(
        "Before you try the rest, read this table. It is what these pictures "
        "actually return on this build, measured on the bench — not what they "
        "ought to return. Two of them do not match their own name, for "
        "reasons that are understood, and a tester who runs them without "
        "knowing that will report a fault that is not there:")
    table(doc, ["Picture", "People in it", "Camera reports", "Why"],
          [["1_person.jpg", "1", "1", "—"],
           ["2_people.jpg", "2", "2", "—"],
           ["3_people.jpg", "3", "3",
            "the picture this step uses; the pass/fail one"],
           ["7_people.jpg", "7", "6 — every time",
            "one of the seven stands behind the others and is mostly hidden. "
            "The network reports what it can see, and it cannot see a whole "
            "person there. Expected; not a fault."],
           ["5_people.jpg", "5", "1 — every time",
            "a wide riverside view where the people are far away and only a "
            "few pixels tall. The network sees the whole scene shrunk to "
            "256x256, and at that size they are not there to be found. A "
            "known limit on distant figures, not a fault in this build."]],
          widths=[1.15, 0.85, 1.05, 3.15])
    note(doc, "So do not judge the step on 7_people or 5_people. Step 7 "
              "passes or fails on the 3-person picture, which is what the "
              "command at the top of this step runs. The other two are here "
              "to show you the edges of what the camera can do, and their "
              "numbers are recorded above so that seeing them does not cost "
              "you a false fault report. If either one gives a different "
              "number from the table, that is worth reporting.", warn=True)
    note(doc, "images/ also has some crowd scenes. The count on a crowd is "
              "approximate by nature — nobody agrees on the true number "
              "either — so use them for a look, never for a judgement.")
    note(doc, "Each extra picture you try loads itself into the camera and "
              "clears itself again the same way, so the last thing any of "
              "these runs does is hand the lens back. Whichever one you "
              "finish on, check its [5/5] line before moving on.")

    # ── Step 8 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 8 — Check an event reaches this PC", level=1)
    doc.add_paragraph(
        "Step 7 proved the camera can count. This proves an event travels the "
        "whole way: camera → internal cable → modem → network → this PC. It "
        "asks the camera to report a detection of 3 people directly, so the "
        "expected numbers are known.")
    doc.add_paragraph("In Window B:")
    cmd(doc, f'{CD}python3 scopus/cam.py "detect simulate 3"')
    expected(doc, "Expected in Window B:")
    code(doc, "[camera /dev/ttyACM2] > detect simulate 3\n"
              "detect simulate: 3 object(s)\n"
              '+SDVRNTF: {"ser":4194336,"num":4,"rsn":16,"rsd":3,'
              '"tim":"20000101002639","mtn":0,\n'
              '           "mod":"","bat":0.0,"vol":0.0}\n'
              "detect simulate ok\n"
              "+SDVRNTF: END,4")
    expected(doc, "Expected in Window A, within a few seconds:")
    code(doc,
         "[15:15:57] NOTIFICATION   from 192.168.2.2  ser=4194336 num=4 "
         "rsn=16 rsd=3\n"
         "[15:15:57]                  valid JSON, 9 fields: "
         '{"ser":4194336,"num":4,"rsn":16,\n'
         '                             "rsd":3,"tim":"20000101002639","mtn":0,'
         '"mod":"","bat":0.0,\n'
         '                             "vol":0.0}')
    doc.add_paragraph("Check three things in Window A, in this order:")
    for a in ["rsn=16 — the reason: 'people detected'.",
              "rsd=3 — how many: three, which is what you asked for.",
              "'valid JSON' — the message arrived intact. This is not "
              "cosmetic. The event is JSON, JSON does not survive the "
              "modem's command channel unaided, and a broken transport shows "
              "up exactly here: a message that arrives but will not read."]:
        doc.add_paragraph(a, style="List Bullet")
    note(doc, "The PASS for this step is in Window A. The camera printing "
              "+SDVRNTF in Window B is not enough — that only proves the "
              "camera spoke, not that anything heard it.")
    note(doc, 'tim="20000101..." means the camera\'s clock is at its factory '
              'default, which happens after it restarts. It does not affect '
              'this test; the event is still real.')

    # ── Step 9 ─────────────────────────────────────────────────────────
    doc.add_heading("Step 9 — The real thing: detect actual people", level=1)
    doc.add_paragraph(
        "This is the product doing its job, with nothing injected or "
        "simulated. Detection is already running from Step 7.")
    doc.add_paragraph(
        "First, one check. This step only works if the camera is looking "
        "through its own lens, and Step 7 is the one thing in this procedure "
        "that takes the lens away from it. Step 7 gives it back on its way "
        "out, so this should already be true — but it costs five seconds to "
        "be sure, and it is the difference between testing the product and "
        "testing a photograph. In Window B:")
    cmd(doc, f'{CD}python3 scopus/cam.py "frame query"')
    expected(doc)
    code(doc, "[camera /dev/ttyACM2] > frame query\n"
              "frame: empty (NN running)\n"
              "frame query ok")
    doc.add_paragraph(
        "'empty' is what you need — no picture loaded, so the camera is on "
        "its lens. 'NN running' means detection is on. If it says "
        "'frame: loaded' instead, the Step 7 picture is still in there and no "
        "amount of walking about will produce an event; the reply tells you "
        "how many seconds until it lapses by itself, and the live view will "
        "be showing TEST PICTURE - NOT THE LENS. Clear it and check again "
        "rather than waiting:")
    cmd(doc, f'{CD}python3 scopus/cam.py "frame clear"')
    note(doc, "If it says 'NN stopped' rather than 'NN running', detection is "
              "off: run  " + CD + "python3 scopus/cam.py \"detect start\"  "
              "and then redo Step 5, which sets what the camera does when it "
              "detects something.")
    doc.add_paragraph()
    doc.add_paragraph("Now the step itself. Do this:")
    for a in ["Make sure nobody is in front of the camera, and wait about "
              "ten seconds. The camera reports the moment people appear "
              "where there were none, so it has to see an empty scene first.",
              "Walk into the camera's view and stay there for a few seconds.",
              "Watch Window A."]:
        doc.add_paragraph(a, style="List Number")
    expected(doc, "Expected in Window A, within a second or two of stepping "
                  "in (this is a real live detection from this bench — three "
                  "people in view):")
    code(doc,
         "[14:53:06] NOTIFICATION   from 192.168.2.2  ser=4194336 num=0 "
         "rsn=16 rsd=3\n"
         "[14:53:06]                  valid JSON, 9 fields: "
         '{"ser":4194336,"num":0,"rsn":16,\n'
         '                             "rsd":3,"tim":"20000101000348","mtn":0,'
         '"mod":"","bat":0.0,\n'
         '                             "vol":0.0}')
    doc.add_paragraph(
        "rsd is the number of people the camera actually sees — 1 for you "
        "alone, 2 if a colleague joins you, 3 in the example above. rsn is "
        "16 again: people detected. Write down how many people were in view "
        "and what rsd said; that comparison is the result of this step.")
    note(doc, "Only the arrival of people is reported, not their continued "
              "presence — standing there does not produce a message every "
              "second, by design. To get another one, step out of view, wait "
              "for the scene to be empty for a few seconds, and step back in.")
    note(doc, "If nothing arrives, work through these three in order. One: "
              "run 'frame query' again — if it says 'loaded', the Step 7 "
              "picture is back in the way and nothing else you try can work. "
              "Two: the camera has to see an empty scene before it can see "
              "people arrive, so if the room already had people in view, step "
              "everyone out, wait ten seconds, then walk back in. Three: if "
              "the live view keeps drawing boxes on an empty room — on a door "
              "frame or a plant, say — the camera believes people are already "
              "there and is waiting for them to leave. Report that with a "
              "photo of the screen; it is a real finding and not something "
              "you can work around.")

    # ── Step 10 ────────────────────────────────────────────────────────
    doc.add_heading("Step 10 — Capture a photo and upload it", level=1)
    doc.add_paragraph(
        "This exercises the other half of the product: a real JPEG from the "
        "lens, across the internal cable, through the modem, onto this PC.")
    cmd(doc, f'{CD}python3 scopus/cam.py "photo upload"')
    expected(doc, "Expected in Window B:")
    code(doc,
         "[camera /dev/ttyACM2] > photo upload\n"
         "photo upload: capturing -> SDVR+SENDBIN ref=2 "
         "name=4194336_01012000_002844.rdy\n"
         '+SDVRNTF: {"ser":4194336,"num":5,"rsn":64,"rsd":2,'
         '"tim":"20000101002844","mtn":0,\n'
         '           "mod":"","bat":0.0,"vol":0.0}\n'
         "photo upload ok\n"
         "+SDVRNTF: END,5\n"
         "+SDVRSRVR: OK")
    doc.add_paragraph(
        "The photo is about 95 KB and crosses the internal cable at 115200 "
        "baud, so give it about 10 seconds, and a second or two more for the "
        "upload itself.")
    expected(doc, "Expected in Window A:")
    code(doc,
         "[15:18:02] NOTIFICATION   from 192.168.2.2  ser=4194336 num=5 "
         "rsn=64 rsd=2\n"
         "[15:18:02]                  valid JSON, 9 fields: {...}\n"
         "[15:18:13] UPLOAD         from 192.168.2.2  93305 bytes  -> "
         "151813_photo\n"
         "[15:18:13]                  JPEG, complete   path=/upload\n"
         "[15:18:13]                  X-Filename: photo\n"
         "[15:18:13]                  X-Filesize: 93305\n"
         "[15:18:13]                  X-Timestamp: 01012000002844\n"
         "[15:18:13]                  X-Ref: 2")
    note(doc, "'JPEG, complete' is the part that matters. It means the file "
              "begins with the JPEG start marker and contains the end marker "
              "— the whole picture arrived. A transfer cut short still "
              "uploads happily and would say TRUNCATED here.")
    note(doc, "The rsn=64 event just above is the camera reporting the "
              "capture itself; it is a different reason code from the rsn=16 "
              "'people detected' events.")

    # ── Step 11 ────────────────────────────────────────────────────────
    doc.add_heading("Step 11 — Look at what actually landed on this PC",
                    level=1)
    doc.add_paragraph("In Window B:")
    cmd(doc, "ls -la ~/scopus-received/")
    expected(doc)
    code(doc, "total 104\n"
              "-rw-rw-r-- 1 user user 93305 Aug  6 15:18 151813_photo\n"
              "-rw-rw-r-- 1 user user   248 Aug  6 15:18 notifications.log")
    doc.add_paragraph()
    cmd(doc, "file ~/scopus-received/*_photo")
    expected(doc)
    code(doc, "/home/user/scopus-received/151813_photo: JPEG image data, "
              "baseline, precision 8, 800x600, components 3")
    doc.add_paragraph()
    cmd(doc, "cat ~/scopus-received/notifications.log")
    expected(doc)
    code(doc,
         '[15:15:57] 192.168.2.2 {"ser":4194336,"num":4,"rsn":16,"rsd":3,'
         '"tim":"20000101002639",\n'
         '                        "mtn":0,"mod":"","bat":0.0,"vol":0.0}\n'
         '[15:18:02] 192.168.2.2 {"ser":4194336,"num":5,"rsn":64,"rsd":2,'
         '"tim":"20000101002844",\n'
         '                        "mtn":0,"mod":"","bat":0.0,"vol":0.0}')
    note(doc, "One line per event: the rsn=16 one is Step 8's people "
              "detection, the rsn=64 one is Step 10's photo capture. Any "
              "events from Step 9 appear here too.")
    doc.add_paragraph()
    doc.add_paragraph("Now open the photo and look at it:")
    cmd(doc, "xdg-open ~/scopus-received/*_photo")
    doc.add_paragraph(
        "It is the scene the camera is pointing at — the real view through "
        "the lens. It is not the test picture from Step 7; that one only "
        "drives the neural network and never becomes a photo.")
    doc.add_paragraph()
    doc.add_paragraph("Check every event arrived exactly once:")
    cmd(doc, "cut -d' ' -f3- ~/scopus-received/notifications.log | sort | "
             "uniq -d")
    expected(doc)
    code(doc, "(no output at all)")
    note(doc, "Any line printed here is the same event delivered twice, "
              "which is a fault worth reporting. Nothing printed is the pass.")

    # ── Step 12 ────────────────────────────────────────────────────────
    doc.add_heading("Step 12 — Finish", level=1)
    doc.add_paragraph("Go to Window A and press Ctrl-C.")
    expected(doc)
    code(doc, "Summary\n"
              "  notifications: 2 received, 2 valid JSON\n"
              "  uploads:       1 received, 1 complete JPEGs\n"
              "  files in /home/user/scopus-received")
    doc.add_paragraph(
        "The two counts on each line must match: every event that arrived "
        "was readable, and every photo that arrived was complete. Record "
        "these numbers, and keep the ~/scopus-received folder — it is the "
        "evidence for this run.")

    # ── Pass criteria ──────────────────────────────────────────────────
    doc.add_heading("13. Pass criteria", level=1)
    doc.add_paragraph("The test passes only if all of these are true:")
    table(doc, ["#", "Must be true", "Where you saw it"],
          [["1", "Every pre-flight check says PASS", "Step 2"],
           ["2", "The modem reads back the endpoints it was given", "Step 4"],
           ["3", "'mdm AT' answers OK — the two boards talk", "Step 6"],
           ["4", "The neural network counts 3 people in the 3-person picture",
            "Step 7"],
           ["5", "An event arrives on this PC with rsn=16 and rsd=3",
            "Step 8, Window A"],
           ["6", "That event is reported 'valid JSON, 9 fields'",
            "Step 8, Window A"],
           ["7", "Real people walking into view produce an event with the "
                 "right count — with 'frame query' saying 'empty', so the "
                 "camera was on its lens and not on the Step 7 picture",
            "Step 9, Window A"],
           ["8", "A photo arrives and is reported 'JPEG, complete'",
            "Step 10, Window A"],
           ["9", "The saved file opens as a picture of the room", "Step 11"],
           ["10", "No event was delivered twice", "Step 11"],
           ["11", "The closing summary shows all events valid and all "
                  "uploads complete", "Step 12"]])
    note(doc, "Testing over the mobile network instead? That is "
              "Scopus_QA_Flow.docx, Part 2, and it carries its own pass "
              "criteria. Say in your report which mode each result came "
              "from — “the event arrived” means something different over the "
              "cable than it does over the mobile network.")

    # ── Troubleshooting ────────────────────────────────────────────────
    doc.add_heading("14. If something fails", level=1)
    table(doc, ["What you see", "What it means and what to do"],
          [["Step 9: people really are in front of the camera, the live view "
            "even draws boxes, but Window A stays silent",
            "The Step 7 picture is almost certainly still loaded, so the "
            "camera is looking at that photograph and not through its lens — "
            "the boxes you see are the photograph's people painted over the "
            "live video, which is why they sit on doors and walls. Check "
            f"with:  {CD}python3 scopus/cam.py \"frame query\"  — if it says "
            f"'loaded', run:  {CD}python3 scopus/cam.py \"frame clear\"  and "
            "redo Step 9. Step 7 clears it for you, so if you find it loaded "
            "here, say so in your report."],
           ["Step 7 says 'the picture did not load', on every attempt",
            "The camera's console and the uploader are out of step — usually "
            "something else was talking to the camera at the same moment. Run "
            "the step once more; it re-syncs on its own. If all four attempts "
            "fail again, check the camera did not restart underneath you:  "
            f"{CD}python3 scopus/cam.py uptime  — a number under a minute "
            "means it did, and that is worth reporting."],
           ["Any command says 'no answer'",
            "Run the same command again — if it was the reply that was lost "
            "rather than the command, the second one answers. This was a "
            "firmware fault fixed on 6 August 2026, so unlike the other rows "
            "here it is worth reporting even when the retry works."],
           ["'ERROR: the modem's AT port was not found', or the modem "
            "answers nothing at all",
            "Almost always another program holding the port — the VS Code "
            "Serial Monitor is the usual one. Close it and run "
            "scopus/preflight.py, which names the program if it is still "
            "there."],
           ["A step worked earlier and now produces no events",
            "The camera has probably restarted, which puts its settings back "
            "to their defaults. Run Step 5's two commands again, then carry "
            "on — and report it. Section 15 explains how to tell, and why an "
            "unprompted restart is now a fault rather than a quirk."],
           ["'mdm AT' does not answer OK",
            "The internal cable between camera and modem is not carrying "
            "traffic. Try it once more, and if it still fails run:  "
            f"{CD}python3 scopus/cam.py \"mdm relink\"  and then 'mdm AT' "
            "again. If that does not fix it, stop and report it — nothing "
            "downstream can work."],
           ["Nothing arrives in Window A at all",
            "Check Window A is still running (it stops if you press Ctrl-C "
            "in it), and run Step 4 again — the modem has to be pointed at "
            "this PC."],
           ["An event arrives but is NOT valid JSON",
            "The message was mangled in transport. Report it with the raw "
            "line Window A prints; this is a real fault."],
           ["'photo upload: trigger failed (busy / no modem)'",
            "A previous capture is still in flight. Wait about 15 seconds "
            "and run it again."],
           ["A photo arrives but says TRUNCATED",
            "The transfer was cut short. Report it, and include the output "
            f"of:  {CD}python3 scopus/cam.py \"mdm stats\""],
           ["The pre-flight says the camera port was not found",
            "The camera re-enumerates after a restart and takes about 30 "
            "seconds to come back. Wait, then run the pre-flight again."]])

    # ── Known behaviour ────────────────────────────────────────────────
    doc.add_heading("15. Known behaviour — not your fault", level=1)
    doc.add_paragraph(
        "The first entry below is normal and needs no action. The second was "
        "a firmware fault, fixed on 6 August 2026 — it is listed because you "
        "may have been told to expect it, and because seeing it now means "
        "something has regressed and should be reported rather than worked "
        "around. The third is still open, and is the one to watch.")
    table(doc, ["Behaviour", "How to recognise it", "What to do"],
          [["An event is slow to arrive",
            "Up to about 10 seconds between the camera reporting and Window "
            "A showing it.",
            "Normal — the camera retries if the modem does not acknowledge. "
            "No action."],
           ["A command's reply goes missing",
            "'no answer to this command', usually right after an event was "
            "sent. The command itself did run.",
            "FIXED — the camera used to stall its own console whenever "
            "nothing was reading the port. Run the command again so you are "
            "not blocked, but report it: it should no longer happen."],
           ["The camera restarts on its own",
            f"Commands stop answering for about 30 seconds; afterwards "
            f"detection is off and the camera's settings are back to their "
            f"defaults. Confirm with:  "
            f"{CD}python3 scopus/cam.py uptime  — a number under a minute "
            f"means it has just restarted.",
            "OPEN — the cause behind the internal cable was fixed, but one "
            "restart was seen on 9 August 2026 during a Step 7 injection and "
            "could not be reproduced afterwards. So this is not closed. Run "
            "'uptime' the moment anything stops answering, redo Step 5 and "
            "carry on, and report it with the time to the second — that "
            "timestamp is what makes it findable. This one matters."]])
    doc.add_paragraph()
    note(doc, "If a restart does happen, the useful evidence is on the "
              "camera's own console at the moment it goes: it prints a line "
              "beginning FAILURE and naming the part that failed, then "
              "restarts about a third of a second later. Nothing captures "
              "that unless a window is watching the camera port at the time. "
              "If you are asked to chase a repeat of this, that is the "
              "recording to make.")
    doc.add_paragraph()
    doc.add_paragraph("If you want to record the state of the link for a "
                      "report:")
    cmd(doc, f'{CD}python3 scopus/cam.py "mdm stats"')
    expected(doc)
    code(doc, "rx: bytes=158 frames=10 badcrc=0 stray=1 err=4 timeouts=861\n"
              "tx: frames=102 err=0 retries=3   usart2 err(ORE/FE/NE)=4\n"
              "ntf: queued=3 sent=3 unconfirmed=0 dropped=0\n"
              "link: relinks=3 consec_timeouts=0")
    table(doc, ["Counter", "What it tells you"],
          [["ntf: queued / sent", "events the camera raised, and events it "
                                  "got to the modem. These should track each "
                                  "other."],
           ["ntf: dropped", "events thrown away. Should be 0 — report it if "
                            "not."],
           ["ntf: unconfirmed", "the modem did not acknowledge in time. It "
                                "does NOT mean the event was lost — judge by "
                                "what reached Window A."],
           ["link: relinks", "times the camera had to reset the internal "
                             "cable. A number that climbs during a test is "
                             "worth reporting."],
           ["rx: badcrc / stray", "corruption on the internal cable. A clean "
                                  "link reports 0 badcrc."]])

    # ── Reporting ──────────────────────────────────────────────────────
    doc.add_heading("16. What to put in your report", level=1)
    for a in ["Pass or fail for each of the eleven criteria in section 13.",
              "The closing summary from Step 12 (the two counts).",
              "For any failure: which step, the exact command you ran, and "
              "everything both windows printed. Copy the text — a "
              "description of it is rarely enough to work out what happened.",
              "How many times you saw anything from section 15, and roughly "
              "when.",
              "The output of the mdm stats command above, taken at the end "
              "of the run.",
              "Keep the ~/scopus-received folder; it holds every event and "
              "photo the test received."]:
        doc.add_paragraph(a, style="List Bullet")

    # ── Automated equivalent ───────────────────────────────────────────
    doc.add_heading("17. The automated version", level=1)
    doc.add_paragraph(
        "The same chain is checked automatically, in far more detail, by the "
        "integration suite. It is not a replacement for this manual — this "
        "one exists to see the product work with your own eyes, and to find "
        "out which hop broke when the suite goes red — but it is the faster "
        "way to check everything at once. Stop the server in Window A first, "
        "since the suite starts its own.")
    cmd(doc, f"{CD}python3 scopus/run_integration_tests.py")
    expected(doc)
    code(doc, "  TOTAL: 66   PASS: 65   FAIL: 0   GAP: 0   SKIP: 1")
    doc.add_paragraph("The one skip is the N6's SD card slot being empty.")

    # ── 18. Remote commands over MQTT ────────────────────────────────
    doc.add_page_break()
    doc.add_heading("18. Driving the unit from the server (remote commands)",
                    level=1)
    doc.add_paragraph(
        "Everything up to here is the unit talking outwards: it sees "
        "something, it reports it, it uploads a photo. This section is the "
        "other direction — typing a command on the server and having the "
        "unit carry it out, over the mobile network, with no cable and "
        "nothing opened on the unit's side.")

    doc.add_heading("Why it works this way", level=2)
    doc.add_paragraph(
        "The unit has no address anyone can dial. On the mobile network it "
        "sits behind the operator's NAT, so the server cannot open a "
        "connection to it. The unit therefore opens one outwards to a "
        "broker and holds it open, and commands are pushed down that "
        "connection.")
    note(doc,
         "This is why the notification channel could not be used for "
         "commands. A reply sent straight back to a report does reach the "
         "unit — measured at 0.2 seconds — but the operator's NAT closes "
         "that opening again in under 30 seconds. Probes sent 30, 60, 120 "
         "and 240 seconds after a report never arrived. So that path can "
         "only answer a report, never start a conversation.")
    doc.add_paragraph(
        "The connection is encrypted and the unit proves who it is with a "
        "certificate. The broker will not accept a client that does not "
        "present one, so there is no password anywhere and no way to drive "
        "someone else's unit.")

    doc.add_heading("What is already set up", level=2)
    table(doc, ["Piece", "Where", "What it is"],
          [["Broker (mosquitto)", f"this PC, port {MQTT_PORT}",
            "Accepts the unit's connection over TLS and requires a client "
            "certificate signed by our CA."],
           ["Certificates", MQTT_CERT_DIR,
            "ca.crt, client.crt, client.key. The same three are already "
            "loaded on the modem under /data/sdvr/certs."],
           ["The unit's name", "its IMEI",
            "Used to build the topics below. Read it with the AT+SDVRMQTT? "
            "command in Step R1 if you do not know it."]])

    doc.add_heading("Step R1 — Check the unit is connected", level=2)
    cmd(doc, f'{CD}python3 scopus/at.py "AT+SDVRMQTT?"')
    expected(doc)
    code(doc, f'+SDVRMQTT: 1,1,"{MQTT_HOST}",{MQTT_PORT},"<IMEI>",0,1,1')
    doc.add_paragraph(
        "The first number is 1 when the channel is switched on, and the "
        "second is 1 when it is actually connected. The text in the middle "
        "is the unit's name — write it down, every command below needs it.")
    note(doc, "If the second number is 0, the unit is not connected. Give it "
              "half a minute and look again: it retries on its own, waiting "
              "5 seconds then 10, 20 and so on up to a minute between "
              "attempts. If it stays 0, check the mobile side first "
              "(Scopus_QA_Flow.docx, Part 2, Step 1) — with no mobile data "
              "there is nothing for it to connect over.")
    doc.add_paragraph("If it is switched off, switch it on with:")
    cmd(doc, f'{CD}python3 scopus/at.py '
             f'"AT+SDVRMQTTSRV=\\"{MQTT_HOST}\\",{MQTT_PORT}"')
    cmd(doc, f'{CD}python3 scopus/at.py "AT+SDVRMQTT=1"')

    doc.add_heading("Step R2 — Watch what the unit says (Window A)", level=2)
    doc.add_paragraph(
        "Open a terminal window and leave this running. It shows every "
        "command that goes out and every answer that comes back. Replace "
        "<IMEI> with the name from Step R1.")
    cmd(doc, f"mosquitto_sub -h {MQTT_HOST} -p {MQTT_PORT} \\\n"
             f"  --cafile {MQTT_CERT_DIR}/ca.crt \\\n"
             f"  --cert {MQTT_CERT_DIR}/client.crt "
             f"--key {MQTT_CERT_DIR}/client.key \\\n"
             f"  -t 'scopus/<IMEI>/#' -v")
    expected(doc, "Straight away, one line:")
    code(doc, "scopus/<IMEI>/status online")
    note(doc, "That line is stored by the broker, so it appears the moment "
              "you subscribe rather than only when the unit next says "
              "something. If it says offline, the unit's connection has "
              "dropped — the broker announces that on the unit's behalf.")

    doc.add_heading("Step R3 — Send a command (Window B)", level=2)
    doc.add_paragraph("In a second window:")
    cmd(doc, f"mosquitto_pub -h {MQTT_HOST} -p {MQTT_PORT} \\\n"
             f"  --cafile {MQTT_CERT_DIR}/ca.crt \\\n"
             f"  --cert {MQTT_CERT_DIR}/client.crt "
             f"--key {MQTT_CERT_DIR}/client.key \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'version'")
    expected(doc, "In Window A, within a few seconds, two lines:")
    code(doc, "scopus/<IMEI>/cmd version\n"
              "scopus/<IMEI>/rsp Application: 01.08.2593089169 "
              "Build: Aug 16 2026 16:24:44 version ok")
    doc.add_paragraph(
        "The first line is your own command coming back past you; the "
        "second is the unit's answer. Anything you can type at the camera "
        "console works here — try uptime, mdm stats, detect start.")

    doc.add_heading("Step R4 — Ask the modem something", level=2)
    doc.add_paragraph(
        "A command beginning with AT is answered by the modem instead of "
        "the camera. This is how a unit in the field gets diagnosed without "
        "anyone travelling to it.")
    cmd(doc, f"mosquitto_pub -h {MQTT_HOST} -p {MQTT_PORT} \\\n"
             f"  --cafile {MQTT_CERT_DIR}/ca.crt \\\n"
             f"  --cert {MQTT_CERT_DIR}/client.crt "
             f"--key {MQTT_CERT_DIR}/client.key \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'AT+CSQ'")
    expected(doc, "In Window A:")
    code(doc, "scopus/<IMEI>/rsp AT+CSQ  +CSQ: 28,99  OK")
    doc.add_paragraph("Useful ones to know:")
    table(doc, ["Command", "What it tells you"],
          [["AT+CSQ", "Signal strength. The first number: under 10 is poor, "
                      "over 20 is good, 99 means no signal at all."],
           ["AT+SDVRNET?", "The whole mobile picture — operator, LTE, APN, "
                           "and the address the unit is using."],
           ["AT+SDVRMQTT?", "This channel's own state and counters."],
           ["AT+SDVRSIM?", "Which SIM holder the unit is reading."]])

    doc.add_heading("Step R5 — The whole product, from the server", level=2)
    doc.add_paragraph(
        "This is the test that matters, because it uses every part at once. "
        "Leave Window A running and send:")
    cmd(doc, f"mosquitto_pub -h {MQTT_HOST} -p {MQTT_PORT} \\\n"
             f"  --cafile {MQTT_CERT_DIR}/ca.crt \\\n"
             f"  --cert {MQTT_CERT_DIR}/client.crt "
             f"--key {MQTT_CERT_DIR}/client.key \\\n"
             f"  -t 'scopus/<IMEI>/cmd' -q 1 -m 'photo upload'")
    expected(doc, "In Window A, an answer like:")
    code(doc, "scopus/<IMEI>/rsp photo upload: capturing -> SDVR+SENDBIN "
              "ref=4 name=4194336_16082026_163709.rdy photo upload ok")
    doc.add_paragraph(
        "Then check the photo itself actually arrived, the same way as "
        "Step 11:")
    cmd(doc, "ls -lt /home/user/sdvr-uploads | head -3")
    expected(doc, "A file called photo, around 100-130 KB, dated just now.")
    note(doc, "A command went from this PC to the unit over the mobile "
              "network, the camera took a picture, and the picture came back "
              "over the same network to this PC. Nothing was plugged in at "
              "either end. That is the whole product in one step.")

    doc.add_heading("Pass criteria for this section", level=2)
    table(doc, ["#", "What must be true"],
          [["R-1", "AT+SDVRMQTT? shows 1,1 — switched on and connected."],
           ["R-2", "The status line says online as soon as you subscribe."],
           ["R-3", "A camera command (version) is answered."],
           ["R-4", "A modem command (AT+CSQ) is answered."],
           ["R-5", "photo upload sent from here puts a photo on this PC."],
           ["R-6", "Every command is answered exactly ONCE. Two identical "
                   "answers to one command is a failure — report it."]])

    doc.add_heading("If the remote commands fail", level=2)
    table(doc, ["What you see", "What it means and what to do"],
          [["AT+SDVRMQTT? shows 1,0 and stays there",
            "Switched on but not connecting. Almost always no mobile data "
            "rather than anything to do with this channel — check the mobile "
            "side (Scopus_QA_Flow.docx, Part 2). If "
            "the mobile side is healthy, the certificates are the next "
            "suspect: the modem must have all three files under "
            "/data/sdvr/certs."],
           ["mosquitto_sub says 'connection refused' or just exits",
            "The broker would not accept your certificate. Check you passed "
            "all three of --cafile, --cert and --key — leaving out --cert is "
            "the usual mistake, and the broker rejects the connection "
            "without explaining why."],
           ["Your command appears on .../cmd but no answer comes back",
            "The command reached the unit and the answer did not come back. "
            "If it was a camera command, the camera link is the suspect — "
            "run Step 6. If it was an AT command, report it."],
           ["The answer arrives twice",
            "A fault, not a quirk. It means a response was published for "
            "each retry instead of once. Report it with the command you "
            "sent."],
           ["'(no output)' comes back",
            "The unit ran something that printed nothing — usually a typing "
            "mistake in the command. Check the spelling against what works "
            "at the console."],
           ["The answer stops mid-sentence with [truncated]",
            "Expected, not a fault: a very long answer is cut after about "
            "1.8 KB. Ask for less at a time."]])

    # ── 20. the motion sensor ────────────────────────────────────────────
    doc.add_heading("19. Checking the motion sensor (the unit being moved)",
                    level=1)
    doc.add_paragraph(
        "This is a different thing from people walking in front of the "
        "camera. The board carries its own movement sensor, and the two "
        "motion notifications — Motion start and Motion stop — mean the "
        "unit itself is being picked up, knocked or tilted. Objects moving "
        "in the picture are reported separately, as People detected and "
        "Vehicle detected.")
    doc.add_paragraph(
        "Two settings control it, and both survive a power cycle:")
    table(doc, ["Setting", "What it does"],
          [["sensitivity 0-100",
            "How hard the box has to be disturbed before it counts. 100 is "
            "the lightest touch (about 15 mg), 0 needs a deliberate shove "
            "(about 500 mg). 50 is the default."],
           ["no-motion timeout",
            "How long the box then has to be still before Motion stop is "
            "sent. Seconds."]])

    doc.add_heading("Step M1 — Check the sensor is there and awake", level=2)
    cmd(doc, f'{CD}python3 scopus/cam.py "motion query"')
    expected(doc)
    code(doc, "motion: sensitivity=50 timeout=30\n"
              "motion: sensor LSM6DSO32 at 0xD6, threshold 265 mg\n"
              "motion: state=still deviation=3 mg peak=0 mg still=142 s\n"
              "motion: events start=0 stop=0")
    cmd(doc, f'{CD}python3 scopus/cam.py "motion read"')
    expected(doc)
    code(doc, "motion read: x=986 y=-77 z=80 mg")
    note(doc,
         "One axis reading about 1000 is gravity. Which axis depends on how "
         "the box is standing. If all three read near zero, the sensor is "
         "answering but not measuring — report it.")

    doc.add_heading("Step M2 — Make it report a movement", level=2)
    doc.add_paragraph(
        "Set a short no-motion timeout first, so you do not wait 30 seconds "
        "for the second event, and allow the two motion events to be "
        "reported:")
    cmd(doc, f'{CD}python3 scopus/cam.py "notify enable 0xff"')
    cmd(doc, f'{CD}python3 scopus/cam.py "motion sense 50 10"')
    doc.add_paragraph("Now move the box: pick it up, put it down, or tap it "
                      "firmly. Watch Window A.")
    expected(doc)
    code(doc, '{"ser":4194336,"num":0,"rsn":2,"rsd":688,'
              '"tim":"20260819170215","mtn":1,…}\n'
              '{"ser":4194336,"num":1,"rsn":4,"rsd":15,'
              '"tim":"20260819170230","mtn":0,…}')
    table(doc, ["Field", "Meaning"],
          [["rsn=2", "Motion start."],
           ["rsn=4", "Motion stop, sent once the box has been still for the "
                     "no-motion timeout."],
           ["rsd", "On start, how big the disturbance was, in mg. On stop, "
                   "how long the movement lasted, in seconds."],
           ["mtn", "1 while the box is moving. It is on every notification, "
                   "not only these two, so an event that arrives during a "
                   "move says so."]])

    doc.add_heading("Step M3 — If nobody can touch the box", level=2)
    doc.add_paragraph(
        "The sensor can push its own measuring element, which produces a "
        "real movement with nobody near the unit. Use this when the bench "
        "is remote:")
    cmd(doc, f'{CD}python3 scopus/cam.py "motion selftest"')
    expected(doc)
    code(doc, "motion selftest: sensor responded (shift 681 mg)")
    doc.add_paragraph(
        "The same two notifications follow, exactly as if the box had been "
        "moved by hand. There is also motion simulate 0|1, which asserts "
        "the state without involving the sensor at all — use it only to "
        "test the reporting path, never to prove the sensor works.")

    doc.add_paragraph("Put the timeout back when you are done:")
    cmd(doc, f'{CD}python3 scopus/cam.py "motion sense 50 30"')

    doc.add_heading("20. When the camera sees it and the count does not",
                    level=1)
    doc.add_paragraph(
        "Sooner or later you will look at the live picture, see a person "
        "plainly, and see the counter say 0. Before reporting it, take one "
        "extra picture: the one the neural network is being given. It is not "
        "the same picture as the one on your screen, and knowing which of "
        "the two is wrong is most of the answer.")
    cmd(doc, f"{CD}python3 n6cam-grab-frame.py --source nn -o nn_input.png")
    expected(doc)
    code(doc,
         "frame grab: nn 256x256 rgb888 196608 bytes\n"
         "frame grab: sensor 2592x1944  nn-area 0,0 2592x1944  "
         "main-area 0,0 2592x1944\n"
         "frame grab: buffer 0x90030000, pipe filling 0x90000000\n"
         "wrote nn_input.png (256x256)")
    doc.add_paragraph(
        "Open nn_input.png. That is exactly what the network saw. Two lines "
        "in the output matter as much as the picture:")
    table(doc,
          ["Line", "What it must say"],
          [["nn-area / main-area",
            "The same rectangle. They are the parts of the sensor the "
            "network and the screen are fed from. If they differ, there is "
            "picture you can see and the network cannot."],
           ["buffer / pipe filling",
            "Two different addresses. If the line says SAME, the picture is "
            "half of one frame and half of the next, and anything moving in "
            "it is torn."]])
    doc.add_paragraph(
        "You can also save what the screen shows, for comparison:")
    cmd(doc, f"{CD}python3 n6cam-grab-frame.py --source live -o live_frame.png")
    note(doc,
         "Attach both PNGs to the report. A missed person is either a "
         "network that was wrong about the picture or a network that was "
         "never shown it, and those have different fixes — these two files "
         "say which, in one step.")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
